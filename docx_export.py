"""
Build a Word (.docx) assembly instruction in the ABICOR BINZEL template layout
from an edited document model produced by the in-browser editor.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config

ASSETS = config.APP_DIR / "assets"
MAGENTA = RGBColor(0xC1, 0x00, 0x6F)
GREY = "BFBFBF"


# ----------------------------------------------------------------- xml helpers
def _shade(paragraph, fill=GREY):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def _thin_borders(table):
    tbl = table._tbl
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "888888")
        borders.append(e)
    tbl.tblPr.append(borders)


def _set_widths(table, widths_mm):
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_mm):
            row.cells[i].width = Mm(w)


def _run(p, text, bold=False, italic=False, size=10, color=None):
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return r


# ----------------------------------------------------------------- builder
def build_docx(model: dict, job_dir: Path, out_path: Path):
    s = model.get("settings", {})
    steps = [st for st in model.get("steps", []) if st.get("include", True)]
    bilingual = s.get("bilingual", False)
    inc_goal = s.get("include_goal", True)
    inc_narr = s.get("include_narration", False)
    inc_parts = s.get("include_part_ids", True)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Mm(14); sec.bottom_margin = Mm(26)
    sec.left_margin = Mm(18); sec.right_margin = Mm(18)

    # ---- header block (title + logo) ----
    htab = doc.add_table(rows=1, cols=2); _no_borders(htab); _set_widths(htab, [120, 54])
    lc, rc = htab.rows[0].cells
    p = lc.paragraphs[0]; _run(p, "BINZEL standard", bold=True, size=20)
    rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo = ASSETS / "logo_header.png"
    if logo.exists():
        rp.add_run().add_picture(str(logo), width=Mm(46))

    _run(doc.add_paragraph(), "Description", size=8)

    bar = doc.add_paragraph(); _shade(bar)
    _run(bar, "Assembly instruction", bold=True, size=12)
    sub = doc.add_paragraph(); _shade(sub)
    _run(sub, "Montageanweisung\n", bold=True, size=8)
    _run(sub, s.get("product_name", "") or s.get("model", ""), bold=True, size=8)

    idp = doc.add_paragraph()
    _run(idp, f"ID-Nummer / ID-Number:  {s.get('id_number','')}", size=9)
    mro = doc.add_paragraph(); _shade(mro); _run(mro, "MRO.", bold=True, size=12)

    st_title = doc.add_paragraph(); st_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(st_title, s.get("station_title", "Station 1: Final Assembly"),
         bold=True, size=13).underline = True
    doc.add_paragraph()

    # ---- steps ----
    for st in steps:
        tab = doc.add_table(rows=1, cols=2); _no_borders(tab); _set_widths(tab, [104, 70])
        tcell, icell = tab.rows[0].cells

        # text cell
        tp = tcell.paragraphs[0]
        _run(tp, f"{st['number']}. {st.get('title','')}", bold=True, size=11)
        if inc_goal and st.get("goal"):
            _run(tcell.add_paragraph(), st["goal"], italic=True, size=8.5,
                 color=RGBColor(0x6b, 0x6b, 0x76))
        for b in st.get("bullets", []):
            bp = tcell.add_paragraph(style="List Bullet")
            _run(bp, b, size=9.5)
        if inc_parts and st.get("parts"):
            for pt in st["parts"]:
                pp = tcell.add_paragraph()
                _run(pp, "■ ", size=8, color=MAGENTA)
                _run(pp, f"{pt.get('part_no','')} ", bold=True, size=8, color=MAGENTA)
                _run(pp, f"– {pt.get('name','')}", size=8, color=MAGENTA)
        if bilingual and st.get("narration_de"):
            _run(tcell.add_paragraph(), st["narration_de"], italic=True, size=8.5)
        if inc_narr and st.get("narration_en"):
            _run(tcell.add_paragraph(), st["narration_en"], size=8.5)

        # image cell
        ip = icell.paragraphs[0]; ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img = _resolve_image(job_dir, st.get("image"))
        if img:
            ip.add_run().add_picture(str(img), width=Mm(66))
            cap = icell.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(cap, f"Step {st['number']}", size=7, color=RGBColor(0x88,0x88,0x88))

        doc.add_paragraph()  # spacer between steps

    _build_footer(doc, s)
    doc.save(str(out_path))
    return out_path


def _resolve_image(job_dir: Path, name):
    if not name:
        return None
    for sub in ("frames", "uploads"):
        p = job_dir / sub / Path(name).name
        if p.exists():
            return p
    p = job_dir / Path(name).name
    return p if p.exists() else None


def _build_footer(doc, s):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    tab = footer.add_table(rows=1, cols=4, width=Mm(174))
    _thin_borders(tab); _set_widths(tab, [70, 38, 38, 28])
    c0, c1, c2, c3 = tab.rows[0].cells
    _run(c0.paragraphs[0], "Drawn by", size=6)
    _run(c0.add_paragraph(), s.get("drawn_by", "AI Documentation Engine"), bold=True, size=7)
    _run(c1.paragraphs[0], "Date of issue", size=6)
    _run(c1.add_paragraph(), s.get("date", ""), bold=True, size=7)
    _run(c2.paragraphs[0], "Document no.", size=6)
    _run(c2.add_paragraph(), s.get("id_number", ""), bold=True, size=7)
    logo = ASSETS / "logo_footer.png"
    if logo.exists():
        c3.paragraphs[0].add_run().add_picture(str(logo), width=Mm(26))
